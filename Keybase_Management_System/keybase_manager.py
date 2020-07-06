from Keybase_Management_System.models import *
from OSINT_System_Core.publisher import publish
from Public_Data_Acquisition_Unit.acquistion_manager import Acquistion_Manager
from django.http import HttpResponse
from Data_Processing_Unit.models import Keybase_Response_TMS

from docx import Document
from docx.shared import Inches



acq = Acquistion_Manager()

class Keybase_Manager(object):


    """
    This class is a Keybase application handler used to access applications models and apply logic


    """



    def __init__(self):
        pass

    def __str__(self):
        pass

    def __repr__(self):
        pass



    #user defined methods bellow

    def create_keybase(self,login_user_id,title,topic,phrases,keywords,mentions,hashtags,**kwargs):
        try:

            kb = Keybase_KMS()
            kb.create(login_user_id,title,topic,phrases,keywords,mentions,hashtags,kwargs)


            publish('keybase created successfully', message_type='control', module_name=__name__)


            return kb

        except Exception as e:
            publish(str(e),message_type='error',module_name=__name__)


    def update_keybase(self,keybase_id,**kwargs):
        try:
            kb = self.get_keybase_object_by_id(keybase_id)

            if 'login_user_id' in kwargs: kb.login_user_id = kwargs['login_user_id']
            if 'title' in kwargs : kb.title = kwargs['title']
            if 'topic' in kwargs: kb.topic = kwargs['topic']
            if 'phrases' in kwargs: kb.phrases = kwargs['phrases']
            if 'keywords' in kwargs: kb.keywords = kwargs['keywords']
            if 'mentions' in kwargs: kb.mentions = kwargs['mentions']
            if 'hashtags' in kwargs: kb.hashtags = kwargs['hashtags']

            kb.save()
            publish('keybase created successfully', message_type='control', module_name=__name__)

            return kb

        except Exception as e:
            publish(str(e),message_type='error',module_name=__name__)

    def get_matched_intell_all(self):
        objects = Keybase_Matched_KMS.objects()
        return objects

    def get_matched_intell_by_date_range(self,start_datetime,end_datetime):
        return Keybase_Matched_KMS.objects(Q(created_on__gte=start_datetime) & Q(created_on__lte=end_datetime))

    def get_matched_intell_by_intell_type(self):
        pass

    def delete_keybase(self,keybase_id):
        try:
            Keybase_KMS.objects(id=keybase_id)[0].delete()
            publish('keybase deleted successfully',message_type='control',module_name=__name__)
        except Exception as e:
            publish(str(e),message_type='error',module_name=__name__)

    def delete_keybase_in_bulk(self,keybase_id_list):
        try:
            for kb_id in keybase_id_list:
                Keybase_KMS.objects(id=kb_id)[0].delete()

            publish('keybases deleted successfully',message_type='control',module_name=__name__)
        except Exception as e:
            publish(str(e),message_type='error',module_name=__name__)

    def get_keybase_object_by_id(self,keybase_id):
        return Keybase_KMS.objects(id=keybase_id)[0]

    def get_keybases_by_date_range(self,start_datetime,end_datetime):
        return Keybase_KMS.objects(Q(created_on__gte=start_datetime) & Q(created_on__lte=end_datetime))

    def get_all_keybases(self):
        return Keybase_KMS.objects()
    
    def get_all_keybases_by_user(self, user_id):
        # req_fields = ['id', 'login_user_id', 'title', 'topic', 'keywords', 'mentions', 'phrases', 'hashtags']
        return Keybase_KMS.objects.filter(login_user_id=user_id).all_fields()
        # .exclude('expire_on').exclude('updated_on').exclude('created_on').exclude('is_enabled').exclude('is_expired')
        # .only('login_user_id', 'title', 'topic', 'keywords', 'mentions', 'phrases', 'hashtags')

    def get_keybase_included_all(self):
        return Keybase_Included_KMS.objects()

    def get_keybase_included_by_date_range(self,start_datetime,end_datetime):
        return Keybase_Included_KMS.objects(Q(created_on__gte=start_datetime) & Q(created_on__lte=end_datetime))

    def get_keybase_processed_report(self,keybase_gtr_id='5f0045ddb244c1312caefd33',records_limit=10):

        #fetch the response object from db

        resp_object = Keybase_Response_TMS.objects(GTR=keybase_gtr_id).first()
        target_object = acq.get_dataobject_by_gtr(acq.get_gtr_by_id(keybase_gtr_id))

        facebook_users = []
        facebook_pages = []
        facebook_groups = []
        facebook_pictures = []
        facebook_videos = []
        facebook_posts = []
        facebook_events = []

        search_engine_acesseble_results = []
        highlighted_instagram_users = []
        highlighted_linkedin_users = []
        highlighted_twitter_users = []


        #filter_empty_user = lambda x:len(x)>0

        resp_object.to_mongo()

        print(target_object)

        for item in resp_object['data']:

            if(item['site'] == 'facebook'):
                for user in item['data']['users'][0:records_limit]:

                    username = user['username']
                    if not len(username) > 0 : username = 'username missing'
                    facebook_users.append((username, user['id'], user['url']))

                for user in item['data']['pages'][0:records_limit]:

                    username = user['username']
                    if not len(username) > 0 : username = 'username missing'
                    facebook_pages.append((username, user['id'], user['url']))

                for user in item['data']['groups'][0:records_limit]:

                    username = user['username']
                    if not len(username) > 0 : username = 'username missing'
                    facebook_groups.append((username, user['id'], user['url']))

                for user in item['data']['posts'][0:records_limit]:

                    username = user['author']['username']
                    if not len(username) > 0 : username = 'username missing'
                    facebook_posts.append((username, user['author']['id'], user['post_id'],user['post_link']))

                # for user in item['data']['pictures'][0:records_limit]:
                #
                #     #username = user['username']
                #     #if not len(username) > 0 : username = 'username missing'
                #
                #     facebook_pictures.append((user['picture_link'], user['referring_url']))
                #
                # for user in item['data']['videos'][0:records_limit]:
                #
                #     username = user['username']
                #     if not len(username) > 0 : username = 'username missing'
                #     facebook_videos.append((username, user['id'], user['url']))

                # for user in item['data']['events'][0:records_limit]:
                #
                #     username = user['username']
                #     if not len(username) > 0 : username = 'username missing'
                #     facebook_users.append((username, user['id'], user['url']))

            if (item['site'] == 'search_engines'):
                for data in item['data']:
                    if(data['status']=='200'):
                        search_engine_acesseble_results.append((data['ip'], data['serp_domain'],data['serp_url']))



            if (item['site'] == 'instagram'):
                for data in item['data']['videos']['users'][0:records_limit]:

                    highlighted_instagram_users.append((data['user']['username'], data['user']['full_name'],data['user']['is_private']))


            if (item['site'] == 'linkedin'):
                for data in item['data']['data'][0:records_limit]:

                    highlighted_linkedin_users.append((data['username'], data['full_name']))


            if (item['site'] == 'twitter'):
                for data in item['data'][0:records_limit]:

                    highlighted_twitter_users.append((data['username'], data['name'],data['user_id_str']))

        #return facebook_users





        document = Document()

        document.add_heading('Report for {0}'.format(target_object.keybase_ref.title), 0)

        p = document.add_paragraph(target_object.keybase_ref.topic)
        #p.add_run('bold').bold = True
        #p.add_run(' and some ')
        #p.add_run('italic.').italic = True

        #document.add_heading('Heading, level 1', level=1)
        #document.add_paragraph('Intense quote', style='Intense Quote')

        #document.add_paragraph(
        #    'first item in unordered list', style='List Bullet'
        #)
        #document.add_paragraph(
        #    'first item in ordered list', style='List Number'
        #)

        # records = (
        #     (3, '101', 'Spam'),
        #     (7, '422', 'Eggs'),
        #     (4, '631', 'Spam, spam, eggs, and spam')
        # )

        if (len(facebook_users) > 0):
            document.add_heading('Facebook Top Matched Users', level=1)
            table = document.add_table(rows=1, cols=3)


            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Usernames'
            hdr_cells[1].text = 'Userids'
            hdr_cells[2].text = 'Urls'

            for username, userid, url in facebook_users:
                row_cells = table.add_row().cells
                row_cells[0].text = username
                row_cells[1].text = userid
                row_cells[2].text = url

        if (len(facebook_pages) > 0):
            document.add_heading('Facebook Top Matched Pages', level=1)
            table = document.add_table(rows=1, cols=3)

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Page Usernames'
            hdr_cells[1].text = 'Page ids'
            hdr_cells[2].text = 'Urls'

            for username, userid, url in facebook_pages:
                row_cells = table.add_row().cells
                row_cells[0].text = username
                row_cells[1].text = userid
                row_cells[2].text = url

        if (len(facebook_groups) > 0):

            document.add_heading('Facebook Top Matched Groups', level=1)
            table = document.add_table(rows=1, cols=3)

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Group Usernames'
            hdr_cells[1].text = 'Group ids'
            hdr_cells[2].text = 'Urls'

            for username, userid, url in facebook_groups:
                row_cells = table.add_row().cells
                row_cells[0].text = username
                row_cells[1].text = userid
                row_cells[2].text = url


        if (len(search_engine_acesseble_results) > 0):
            document.add_heading('Highlighted Accessible Sites', level=1)
            table = document.add_table(rows=1, cols=3)

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'IP Address'
            hdr_cells[1].text = 'Domain'
            hdr_cells[2].text = 'Url'

            for username, userid, url in search_engine_acesseble_results:
                row_cells = table.add_row().cells
                row_cells[0].text = username
                row_cells[1].text = userid
                row_cells[2].text = url

        if (len(highlighted_instagram_users) > 0):

            document.add_heading('Highlighted Instagram Users', level=1)
            table = document.add_table(rows=1, cols=3)

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Usernames'
            hdr_cells[1].text = 'Full Name'
            hdr_cells[2].text = 'Private'

            for username, userid, url in highlighted_instagram_users:
                row_cells = table.add_row().cells
                row_cells[0].text = username
                row_cells[1].text = userid
                row_cells[2].text = str(url)

        if (len(highlighted_linkedin_users) > 0):
            document.add_heading('Highlighted Linkedin Users', level=1)
            table = document.add_table(rows=1, cols=3)

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Usernames'
            hdr_cells[1].text = 'Full Name'


            for username, userid in highlighted_linkedin_users:
                row_cells = table.add_row().cells
                row_cells[0].text = username
                row_cells[1].text = userid


        if(len(highlighted_twitter_users)>0):
            document.add_heading('Highlighted Twitter Users', level=1)
            table = document.add_table(rows=1, cols=3)

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Usernames'
            hdr_cells[1].text = 'Full Name'
            hdr_cells[2].text = 'Private'

            for username, userid, url in highlighted_twitter_users:
                row_cells = table.add_row().cells
                row_cells[0].text = username
                row_cells[1].text = userid
                row_cells[2].text = str(url)

        document.add_page_break()

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=download.docx'
        document.save(response)
        return response